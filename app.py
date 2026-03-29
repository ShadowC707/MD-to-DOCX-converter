import streamlit as st
import subprocess
import tempfile
import os
import zipfile
import glob
import re
from docx import Document
from docx.shared import Pt

st.set_page_config(page_title="MD → DOCX", page_icon="📄")

st.title("MD → DOCX")
st.caption("Upload a Markdown file and get a Word document back. Formulas and images included.")

FONT_OPTIONS = [
    "Times New Roman",
    "Arial",
    "Calibri",
    "Georgia",
    "Verdana",
]

SIZE_OPTIONS = [12, 13, 14, 16]

SPACING_OPTIONS = {
    "Single (1.0)": 1.0,
    "1.5 lines": 1.5,
    "Double (2.0)": 2.0,
}

def fix_image_syntax(text):
    return re.sub(r'!\[\[(.+?)\]\]', r'![](\1)', text)

def apply_font(docx_path, font_name, font_size, line_spacing):
    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
        paragraph.paragraph_format.line_spacing = line_spacing
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
    doc.save(docx_path)

# ---- Formatting options (always visible) ----
col1, col2, col3 = st.columns(3)
with col1:
    font_name = st.selectbox("Font", FONT_OPTIONS)
with col2:
    font_size = st.selectbox("Font size", SIZE_OPTIONS, index=2)
with col3:
    spacing_label = st.selectbox("Line spacing", list(SPACING_OPTIONS.keys()), index=1)
    line_spacing = SPACING_OPTIONS[spacing_label]

fix_obsidian = st.checkbox("Fix Obsidian image syntax (convert ![[image]] embeds)")

st.divider()

# ---- Input tabs ----
tab_upload, tab_paste = st.tabs(["Upload file", "Paste text"])

uploaded_file = None
pasted_text = None
output_filename = "converted.docx"

with tab_upload:
    st.info("""
**Two ways to upload:**
- **Single `.md` file** — no images
- **`.zip` file** — your `.md` + any images in the same folder
""")
    uploaded_file = st.file_uploader("Choose a .md or .zip file", type=["md", "zip"])

with tab_paste:
    pasted_text = st.text_area(
        "Paste your Markdown here",
        height=300,
        placeholder="# My document\n\nHere is a formula: $E = mc^2$\n\nAnd some **bold text**."
    )

# ---- Convert button ----
st.divider()

has_input = (uploaded_file is not None) or (pasted_text and pasted_text.strip())

if st.button("Convert", type="primary", disabled=not has_input):
    with st.spinner("Converting..."):
        with tempfile.TemporaryDirectory() as tmpdir:

            input_path = os.path.join(tmpdir, "input.md")

            if uploaded_file is not None:
                # --- File upload branch ---
                if uploaded_file.name.endswith(".zip"):
                    zip_path = os.path.join(tmpdir, "upload.zip")
                    with open(zip_path, "wb") as f:
                        f.write(uploaded_file.getvalue())

                    with zipfile.ZipFile(zip_path, "r") as z:
                        z.extractall(tmpdir)

                    md_files = glob.glob(os.path.join(tmpdir, "**", "*.md"), recursive=True)
                    if not md_files:
                        st.error("No .md file found inside the zip.")
                        st.stop()

                    input_path = md_files[0]
                    output_filename = os.path.basename(input_path).replace(".md", ".docx")

                    if fix_obsidian:
                        with open(input_path, "r", encoding="utf-8") as f:
                            content = f.read()
                        with open(input_path, "w", encoding="utf-8") as f:
                            f.write(fix_image_syntax(content))

                else:
                    content = uploaded_file.getvalue().decode("utf-8")
                    if fix_obsidian:
                        content = fix_image_syntax(content)
                    with open(input_path, "w", encoding="utf-8") as f:
                        f.write(content)
                    output_filename = uploaded_file.name.replace(".md", ".docx")

            else:
                # --- Paste branch ---
                content = pasted_text
                if fix_obsidian:
                    content = fix_image_syntax(content)
                with open(input_path, "w", encoding="utf-8") as f:
                    f.write(content)
                output_filename = "converted.docx"

            output_path = os.path.join(tmpdir, output_filename)

            result = subprocess.run(
                ["pandoc", input_path, "-o", output_path, "--mathml"],
                capture_output=True,
                text=True,
                cwd=os.path.dirname(input_path)
            )

            if result.returncode != 0:
                st.error(f"Pandoc error: {result.stderr}")
            elif not os.path.exists(output_path):
                st.error("Conversion failed — no output file produced.")
            else:
                apply_font(output_path, font_name, font_size, line_spacing)

                with open(output_path, "rb") as f:
                    docx_bytes = f.read()

                st.success("Done!")
                st.download_button(
                    label="⬇ Download .docx",
                    data=docx_bytes,
                    file_name=output_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )